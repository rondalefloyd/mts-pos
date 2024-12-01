import os, sys, logging, math, json, pandas as pd, win32com.client, pythoncom
from peewee import *
from datetime import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from docx import Document

sys.path.append(os.path.abspath(''))  # required to change the default path
from app.models.entities import *
from app.utils.databases import postgres_db


class PrintThread(QThread):
    running = pyqtSignal(object)
    finished = pyqtSignal(object)
    
    def __init__(self, function_route, entry=None):
        super().__init__()
        self.function_route = function_route
        self.entry = entry
        self.isActive = True
    
    def run(self):
        result = {
            'success': False,
            'message': 'N/A',
            'dictData': {},
            'listData': [],
        }
         
        try:
            with postgres_db:
                if self.function_route == 'printReceipt':
                    result = self.printReceipt(self.entry, result)
                else:
                    result['message'] = f"'{self.function_route}' is an invalid function..."
                        
            logging.info('database operation done...')
            
        except Exception as exception:
            result['message'] = f"An error occured: {exception}"
            postgres_db.rollback()
            logging.error('exception: %s', exception)
            logging.info('database operation rolled back...')
            
        finally:
            postgres_db.close()
            logging.info('database closed...')
            
        self.finished.emit(result)
        print(f"{self.function_route} -> result_message: {result['message']}")

    def stop(self):
        self.isActive = False  # Set the flag to stop the thread
        
    # add function here
    def printReceipt(self, entry=None, result=None):
        try:
            pythoncom.CoInitialize()
            # Load the DOCX template
            template_path = os.path.abspath('app/utils/receipt_template.docx')
            document = Document(template_path)
            
            # Define the mapping of placeholders to their corresponding values from the JSON object
            organization = Organization.get_or_none(Organization.Id == entry['organizationId'])
            user = User.get_or_none(User.Id == entry['userId'])
            order = entry['order']
            billing = entry['billing']
            cart = order['cart']
            
            placeholders = {
                '<OrganizationName>': organization.OrganizationName,
                '<Address>': organization.Address,
                '<TransactionDate>': datetime.now().strftime('%Y-%m-%d'),  # Example: current date
                '<ReferenceId>': order['referenceId'],
                '<TaxId>': organization.TaxId,
                '<MachineId>': order['machineId'],  # Example: static value
                '<Quantity>': '\n'.join([str(item['quantity']) for item in cart]),
                '<ItemName>': '\n'.join([item['itemName'] for item in cart]),
                '<Total>': '\n'.join([str(item['total']) for item in cart]),
                '<Subtotal>': str(billing['subtotal']),
                '<Discount>': str(billing['discount']),
                '<Tax>': str(billing['tax']),
                '<GrandTotal>': str(billing['grandtotal']),
                '<Payment>': str(billing['payment']),
                '<PaymentType>': billing['paymentType'],
                '<Change>': str(billing['change']),
                '<UserName>': user.UserName,
                '<MobileNumber>': user.MobileNumber,
            }
            
            result['dictData'] = {
                'dataRepresentation': None,
                'currentDataCount': 0,
                'totalDataCount': len(placeholders),
            }
            dictData = result['dictData']
            
            # Replace placeholders in paragraphs and tables
            elements = document.paragraphs + [paragraph for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]
            for element in elements:
                for run in element.runs:
                    for placeholder, value in placeholders.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            
                            dictData['dataRepresentation'] = value
                            dictData['currentDataCount'] += 1
                            self.running.emit(dictData)
            
            # Save the modified document
            output_path = os.path.abspath('app/utils/output.docx')
            document.save(output_path)
            
            # Print the document (Windows-specific)
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(output_path)
            doc.PrintOut()
            doc.Close(False)
            word.Quit()
            
            pythoncom.CoInitialize()
            
            result['success'] = True
            result['message'] = 'Receipt generated and sent to printer successfully'
            result['dictData'] = placeholders  # Optionally return the placeholders and their values
            return result

        except Exception as exception:
            result['success'] = False
            result['message'] = f"An error occured: {exception}"
            return result
        
        